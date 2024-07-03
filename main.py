from docx import Document
import os
from groq import Groq
import sys

client = Groq(
    api_key=os.environ.get("GROQ_API_KEY"),
)

#check if file path is provided as CLA
if len(sys.argv) < 2:
    print("Error: No file path provided.")
    sys.exit(1)

file_path = sys.argv[1]
document = Document(file_path)

prev_edit = None
for para in document.paragraphs:
    if prev_edit:
        para.insert_paragraph_before(prev_edit, style='ListBullet')
    text = para.text
    edited_text = client.chat.completions.create(
    messages=[
        {
            "role": "system",
            "content": "You're a professional editor in academia and your're editing a very important research paper.\n\nRe-write each paragraph you're given with appropriate word choices and sentence structures, only change where necessary.  Make sure your output is semantically identical.  Be clear, concise, and spartan.  You only have one try, so double check your edit for accuracy and readability.  Only output the edited paragraph.",
        },
        {
            "role": "user",
            "content": text
        }
    ],
    model="llama3-8b-8192",
    )
    prev_edit = edited_text.choices[0].message.content

document.add_paragraph(prev_edit, style='ListBullet')

document.save(file_path)